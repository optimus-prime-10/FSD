import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Helper to add code blocks
def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.style = 'Normal'
    return p

def main():
    doc = Document()
    doc.add_heading('Multipage Websites Learning Guide', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('This guide covers all aspects of building multipage websites, including file paths, project structure, HTML boilerplate, portfolio projects, and hosting. Each section combines theory and practical code examples for easy teaching.\n')

    base = '2. Multipage Websites'
    root = os.path.abspath(base)

    # 1. Computer File Paths
    doc.add_heading('1. Computer File Paths', level=1)
    with open(os.path.join(root, '1. Computer file paths/summary.txt')) as f:
        doc.add_paragraph(f.read())
    doc.add_paragraph('Example: Folder0/index.html (before):')
    with open(os.path.join(root, '1. Computer file paths/Folder0/index.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Solution: Folder0/solution.html (with correct file paths):')
    with open(os.path.join(root, '1. Computer file paths/Folder0/solution.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Note: Images referenced in the code (e.g., cat.png, dog.png) would appear here if this were a web page.')

    # 2. Webpages and Multi-Page Websites
    doc.add_page_break()
    doc.add_heading('2. Webpages and Multi-Page Websites', level=1)
    with open(os.path.join(root, '2. Webpages/summary.txt')) as f:
        doc.add_paragraph(f.read())
    doc.add_paragraph('index.html:')
    with open(os.path.join(root, '2. Webpages/index.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Solution (solution.html):')
    with open(os.path.join(root, '2. Webpages/solution.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('About page (about.html):')
    with open(os.path.join(root, '2. Webpages/public/about.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Contact page (contact.html):')
    with open(os.path.join(root, '2. Webpages/public/contact.html')) as f:
        add_code_block(doc, f.read())

    # 3. HTML Boilerplate
    doc.add_page_break()
    doc.add_heading('3. HTML Boilerplate', level=1)
    with open(os.path.join(root, '3. HTML boilderplate/summary.txt')) as f:
        doc.add_paragraph(f.read())

    # 4. Portfolio Website Project
    doc.add_page_break()
    doc.add_heading('4. Portfolio Website Project', level=1)
    with open(os.path.join(root, '4. Portfolio Website/summary.txt')) as f:
        doc.add_paragraph(f.read())
    doc.add_paragraph('index.html (starter with TODOs):')
    with open(os.path.join(root, '4. Portfolio Website/index.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Solution (solution.html):')
    with open(os.path.join(root, '4. Portfolio Website/solution.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('About page (about.html):')
    with open(os.path.join(root, '4. Portfolio Website/public/about.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Contact page (contact.html):')
    with open(os.path.join(root, '4. Portfolio Website/public/contact.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Movie Ranking Project (movie-ranking.html):')
    with open(os.path.join(root, '4. Portfolio Website/public/movie-ranking.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Birthday Invite Project (birthday-invite.html):')
    with open(os.path.join(root, '4. Portfolio Website/public/birthday-invite.html')) as f:
        add_code_block(doc, f.read())
    doc.add_paragraph('Note: Images referenced in the code (e.g., movie-ranking.png, birthday-invite.png) would appear here if this were a web page.')

    # 5. Hosting on Github
    doc.add_page_break()
    doc.add_heading('5. Hosting on Github', level=1)
    with open(os.path.join(root, '5. Hosting on Github/summary.txt')) as f:
        doc.add_paragraph(f.read())

    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('You have now learned how to build, structure, and host multipage websites. Practice by building your own projects and sharing them online!')

    doc.save('Multipage_Websites_Learning_Guide.docx')
    print('Multipage_Websites_Learning_Guide.docx created successfully.')

if __name__ == '__main__':
    main() 